import docx

doc = docx.Document()
doc.add_heading('John Doe', 0)
doc.add_paragraph('Python Developer with 5 years of experience in Django and Machine Learning.')
doc.add_heading('Skills', level=1)
doc.add_paragraph('Python, Django, SQL, Machine Learning, React, AWS')
doc.add_heading('Experience', level=1)
doc.add_paragraph('Senior Developer at Tech Corp. Built AI resume screeners.')
doc.save('dummy_resume.docx')
print("Dummy resume created.")
